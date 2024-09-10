from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from bs4 import BeautifulSoup
import time
from difflib import SequenceMatcher
from pathlib import Path
import glob
from selenium.webdriver import ActionChains
from selenium.webdriver.common.actions.action_builder import ActionBuilder
from pydub import AudioSegment
from pydub.silence import split_on_silence
import os
import random
import json
import re
from docx import Document

def filter_items_by_string(input_list, target_string):
    filtered_list = [item for item in input_list if target_string in item]
    return filtered_list

chrome_options = webdriver.ChromeOptions()
download_dir = os.path.abspath('./audiosTnu')
os.makedirs(download_dir, exist_ok=True)
prefs = {'download.default_directory' : download_dir}
chrome_options.add_experimental_option('prefs', prefs)

servico = Service(ChromeDriverManager().install())

navegador = webdriver.Chrome(service=servico, options=chrome_options)

navegador.get("https://www.cjf.jus.br/jurisprudencia/tnu/")

#Usando o selenium para clicar no botão pesquisa avançada

navegador.find_element('xpath', '//*[@id="formulario:ckbAvancada"]/div[2]/span').click()

#Usando time para fazer o Selenium esperar carregar página do navegador

time.sleep(1)

#Clicando e preenchendo o fomulário com as datas que queremos

navegador.find_element('xpath', '//*[@id="formulario:j_idt25_input"]').send_keys("01/07/2022")

navegador.find_element('xpath', '//*[@id="formulario:j_idt27_input"]').send_keys("01/08/2023")

#Clicando no botão enviar, para efetuar a pesquisas

navegador.find_element('xpath', '//*[@id="formulario:actPesquisar"]/span').click()

time.sleep(5)

soup = BeautifulSoup(navegador.page_source, "html.parser")

#Criando a lista de links do interior teor dos acordãos que queremos

ementasLista = []
numberList = []
rapporteurList = []
time.sleep(5)

time.sleep(5)

quantity = 0

document = Document()

ementas = soup.findAll('div', id=re.compile('^painel_ementa-'))
for ementa in ementas:
    ementasLista.append(ementa)
    quantity = len(ementasLista)
    xpath_expression1 = f"/html/body/div[3]/form/div/div[3]/div[3]/div[1]/div[2]/div[{quantity}]/div/table/tbody/tr[2]/td/div/div/div/div/div/table/tbody/tr[4]/td"
    xpath_expression2 = f"/html/body/div[3]/form/div/div[3]/div[3]/div[1]/div[2]/div[{quantity}]/div/table/tbody/tr[2]/td/div/div/div/div/div/table/tbody/tr[8]/td"
    elementProcessNumber = navegador.find_element('xpath', xpath_expression1)
    elementTNURapporteur = navegador.find_element('xpath', xpath_expression2)
    numberList.append(elementProcessNumber.text)
    rapporteurList.append(elementTNURapporteur.text)
    #document.add_paragraph(elementTNURapporteur.text)  # Add rapporteur to the Word document
    #document.add_paragraph(elementProcessNumber.text)  # Add process number to the Word document
    #document.add_paragraph(ementa.text)  # Add the text to the Word document
    print(len(ementasLista))

while True:
    try:
        navegador.find_element('xpath', '//*[@id=\"formulario:tabelaDocumentos_paginator_top\"]/a[3]/span').click()
        time.sleep(5)
        soup = BeautifulSoup(navegador.page_source, "html.parser")
        ementas = soup.findAll('div', id=re.compile('^painel_ementa-'))
        for index, ementa in enumerate(ementas):
            ementasLista.append(ementa.text)
            quantity = (index % 30) + 1  
            xpath_expression1 = f"/html/body/div[3]/form/div/div[3]/div[3]/div[1]/div[2]/div[{quantity}]/div/table/tbody/tr[2]/td/div/div/div/div/div/table/tbody/tr[4]/td"
            xpath_expression2 = f"/html/body/div[3]/form/div/div[3]/div[3]/div[1]/div[2]/div[{quantity}]/div/table/tbody/tr[2]/td/div/div/div/div/div/table/tbody/tr[8]/td"
            elementProcessNumber = navegador.find_element('xpath', xpath_expression1)
            elementTNURapporteur = navegador.find_element('xpath', xpath_expression2)
            numberList.append(elementProcessNumber.text)
            rapporteurList.append(elementTNURapporteur.text)
            #document.add_paragraph(elementTNURapporteur.text)  # Add rapporteur to the Word document
            #document.add_paragraph(elementProcessNumber.text)  # Add process number to the Word document
            #document.add_paragraph(ementa.text)  # Add the text to the Word document

    except Exception as e:
        print("Erro: ", e)
        break

original_list = ementasLista
string_to_check = "UNIFORMIZAÇÃO"
filtered_list = filter_items_by_string(original_list, string_to_check)
print(filtered_list)

for item in filtered_list:
    document.add_paragraph(item)

# Salvando em word
document.save("ementas.docx")

time.sleep(5)

